from selenium import webdriver
import docx
# 封装写入文档的方法
def write_docx(title,data):
    word = docx.Document()
    word.add_heading(title,0)
    word.add_paragraph(str(data))
    path = "D:\python\webauto\\"+title+".docx"
    path_text = "D:\python\webauto\\"+title+".txt"
    word.save(path)
    with open(path_text,"w",encoding="utf-8") as file:
        file.write(str(data))
def get_data(method_name,css_selector_fornt):
    driver = webdriver.Firefox()
    driver.get("http://doc.redisfans.com/")
    driver.maximize_window()
    driver.implicitly_wait(20)
    method_name = []
    methods_url = {}
    methods_ex = {}
    data = {}
    for i in range(1,100):
        driver.get("http://doc.redisfans.com/")
        i = str(i)
        css = css_selector_fornt+i+") > a:nth-child(1)"
        try:
            method_name.append(driver.find_element_by_css_selector(css).get_attribute("textContent"))
            driver.find_element_by_css_selector(css).click()
            methods_url[method_name[-1]] = driver.current_url
            methods_ex[method_name[-1]] = driver.find_element_by_css_selector(".highlight-python > pre:nth-child(1)").text
            data[method_name[-1]] = [{methods_ex[method_name[-1]]}, {methods_url[method_name[-1]]}]
        except:
            print("已完成所有方法点击")
            break
    driver.quit()
    return data
def run(title,css_front,):
    data = get_data(title,css_front)
    write_docx(title,data)
if __name__ == '__main__':
    css_key_fornt = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(1) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_string_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(2) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_hash_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(3) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_list_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(1) > td:nth-child(4) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_set_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(2) > td:nth-child(1) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_sorted_set_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(2) > td:nth-child(2) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_pub_sub_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(2) > td:nth-child(3) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_Transaction_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(2) > td:nth-child(4) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_script_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(3) > td:nth-child(1) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_connection_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(3) > td:nth-child(2) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    css_server_front = "div.wy-table-responsive:nth-child(2) > table:nth-child(1) > tbody:nth-child(2) > tr:nth-child(3) > td:nth-child(3) > div:nth-child(1) > ul:nth-child(1) > li:nth-child(1) > ul:nth-child(2) > li:nth-child("
    run("key",css_key_fornt)
    run("string",css_string_front)
    run("hash",css_hash_front)
    run("list",css_list_front)
    run("set",css_set_front)
    run("sorted_set",css_sorted_set_front)
    run("pub_sub",css_pub_sub_front)
    run("transaction",css_Transaction_front)
    run("script",css_script_front)
    run("connection",css_connection_front)
    run("server",css_server_front)